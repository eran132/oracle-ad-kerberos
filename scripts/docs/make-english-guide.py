# -*- coding: utf-8 -*-
"""Generate the English LTR .docx build guide (mirror of BUILD-STEPS)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(11)

def _mk(tag, **a):
    e = OxmlElement(tag)
    for k, v in a.items():
        e.set(qn(k), v)
    return e

def heading(text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = BODY_FONT
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
    return p

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold
    r.font.name = BODY_FONT; r.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def step(tag, num, text):
    p = doc.add_paragraph()
    r0 = p.add_run(f"{num}. "); r0.bold = True; r0.font.name = BODY_FONT
    rt = p.add_run(f"[{tag}] "); rt.bold = True; rt.font.name = MONO_FONT
    rt.font.color.rgb = RGBColor(0xB0, 0x30, 0x10)
    rb = p.add_run(text); rb.font.name = BODY_FONT; rb.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    return p

def success(text):
    p = doc.add_paragraph()
    r0 = p.add_run("✔ Success: "); r0.bold = True; r0.font.name = BODY_FONT
    r0.font.color.rgb = RGBColor(0x1B, 0x7A, 0x1B)
    r1 = p.add_run(text); r1.font.name = BODY_FONT; r1.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(8)
    return p

def note(text):
    p = doc.add_paragraph()
    r0 = p.add_run("⚠ Note: "); r0.bold = True; r0.font.name = BODY_FONT
    r0.font.color.rgb = RGBColor(0xB0, 0x60, 0x00)
    r1 = p.add_run(text); r1.font.name = BODY_FONT; r1.italic = True; r1.font.size = Pt(10.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def code(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3); pf.right_indent = Inches(0.3)
    pf.space_before = Pt(2); pf.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pPr.append(_mk("w:shd", **{"w:val": "clear", "w:fill": "F2F2F2"}))
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line); r.font.name = MONO_FONT; r.font.size = Pt(9.5)
    return p

# ===========================================================================
heading("Setup Guide: Active Directory authentication to Oracle via Kerberos", 0)
para("This document is a step-by-step operational thread for what to do on each component so a domain (Active Directory) user can connect to Oracle through DBeaver on Windows using Kerberos, with no local Oracle password. Each step is tagged with the machine it runs on: [AD] = the Domain Controller, [ORA] = the Oracle server, [WKS] = the Windows workstation.")
note("All commands, product names and file paths are literal — copy them as-is. Lab placeholders to search-replace for a real environment: MYLAB.LOCAL / mylab.local, ad1.mylab.local, ora01.mylab.local, wks01.mylab.local, groups oracle-readers / oracle-writers, test users alice / bob / carol.")

heading("Phase 0 — Prerequisites", 1)
step("AD", 1, "Windows Server 2022 DC promoted, forest mylab.local, DNS role installed, A-records for ad1, ora01, wks01.")
step("AD", 2, "AD Certificate Services (Enterprise Root CA named mylab-root-ca) installed. Confirm the DC auto-enrolled an LDAPS cert (subject CN=ad1.mylab.local).")
step("ORA", 3, "RHEL 9 host with Oracle 19c (19.30+), CDB ORCLCDB open, PDB ORCLPDB1 open READ WRITE, Listener on port 1521.")
step("ALL", 4, "All three hosts on the same routable network. Verify reachability to ports 88, 389, 636, 464.")
step("ALL", 5, "Clock skew between every pair of hosts < 300 seconds (Kerberos window). ora01 syncs to ad1 via chronyd.")

heading("Phase 1 — Active Directory setup (on ad1)", 1)
step("AD", 6, "Create the svc-ora01 service account with a THROWAWAY password (placeholder only — AD won't create an enabled account without one; step 7's ktpass resets it to the value that counts). If it already exists: skip New-ADUser, just Enable-ADAccount + Set-ADUser.")
code("$throwaway = ConvertTo-SecureString ([guid]::NewGuid().ToString()+'Aa1!') -AsPlainText -Force\nNew-ADUser -Name svc-ora01 -SamAccountName svc-ora01 `\n  -UserPrincipalName svc-ora01@MYLAB.LOCAL -AccountPassword $throwaway `\n  -Enabled $true -PasswordNeverExpires $true -CannotChangePassword $true `\n  -Path \"CN=Users,DC=mylab,DC=local\"\nSet-ADUser -Identity svc-ora01 -KerberosEncryptionType \"AES256\"")
success("Get-ADUser svc-ora01 returns; msDS-SupportedEncryptionTypes = 16.")
step("AD", 7, "Register the SPN and emit the keytab in one shot with ktpass. Use -pass * so it prompts (hidden). This is the AUTHORITATIVE password — store it in your password vault. ktpass ALWAYS resets the password and bumps KVNO, so running this command is identical for first-time setup and rotation.")
code("ktpass -princ oracle/ora01.mylab.local@MYLAB.LOCAL `\n       -mapuser MYLAB\\svc-ora01 `\n       -pass * `\n       -crypto AES256-SHA1 -ptype KRB5_NT_PRINCIPAL `\n       -out C:\\temp\\ora01.keytab")
note("The two flags people get wrong: -crypto AES256-SHA1 = derive ONLY the AES256 (etype 18) key into the keytab; must match the account's msDS-SupportedEncryptionTypes (=16), sqlnet.ora and the client krb5.ini; never use All (ships a weak RC4 key). -ptype KRB5_NT_PRINCIPAL = standards-compliant principal name-type Oracle/MIT/Java expect; other types cause name-type mismatches that fail auth. Full table + AES verification: docs/19-ad-admin-runbook.md section A2. After generating, verify klist -kte shows etype 18 and Get-ADUser svc-ora01 -Properties msDS-SupportedEncryptionTypes returns 16.")
note("If setspn -Q oracle/ora01.mylab.local returns more than one account or one that is NOT svc-ora01, remove the SPN from each wrong holder: setspn -D oracle/ora01.mylab.local <wrong-account>, then re-run ktpass.")
success("setspn -Q oracle/ora01.mylab.local returns exactly one account (CN=svc-ora01). setspn -X reports zero duplicates.")
step("AD", 8, "Create the svc-ora-ldap service account (Oracle binds to AD over LDAPS with it for group lookups). Default Domain Users privileges suffice. Unlike svc-ora01 it has NO keytab — its password is loaded directly into the Oracle wallet, so the password set here IS authoritative. If it already exists: do NOT change the password (that breaks the sync; rotation is the deliberate Part B3 procedure).")
code("if (-not (Get-ADUser -Filter \"SamAccountName -eq 'svc-ora-ldap'\" `\n          -ErrorAction SilentlyContinue)) {\n  $pw = Read-Host -AsSecureString \"NEW password for svc-ora-ldap\"\n  New-ADUser -Name svc-ora-ldap -SamAccountName svc-ora-ldap `\n    -UserPrincipalName svc-ora-ldap@MYLAB.LOCAL -AccountPassword $pw `\n    -Enabled $true -PasswordNeverExpires $true -CannotChangePassword $true `\n    -Path \"CN=Users,DC=mylab,DC=local\"\n} else { Write-Host \"exists - leave password (rotate via Part B3)\" }")
success("Get-ADUser svc-ora-ldap returns enabled.")
step("AD", 9, "Create the OU and groups: OU=Groups,DC=mylab,DC=local containing oracle-readers and oracle-writers (Global Security groups). Idempotent — safe to re-run.")
code("if (-not (Get-ADOrganizationalUnit -Filter 'Name -eq \"Groups\"' `\n          -SearchBase \"DC=mylab,DC=local\" -ErrorAction SilentlyContinue)) {\n  New-ADOrganizationalUnit -Name \"Groups\" -Path \"DC=mylab,DC=local\"\n}\nNew-ADGroup -Name oracle-readers -GroupScope Global -GroupCategory Security `\n  -Path \"OU=Groups,DC=mylab,DC=local\"\nNew-ADGroup -Name oracle-writers -GroupScope Global -GroupCategory Security `\n  -Path \"OU=Groups,DC=mylab,DC=local\"")
step("AD", 10, "Add the test users to the groups (alice + carol to oracle-readers, bob to oracle-writers).")
code("Add-ADGroupMember -Identity oracle-readers -Members alice, carol\nAdd-ADGroupMember -Identity oracle-writers -Members bob")
step("AD", 11, "Export the Root CA cert so the Oracle wallet can trust the LDAPS chain.")
code("(Get-CACertificate).RawData | Set-Content -Encoding Byte C:\\temp\\mylab-root-ca.cer\ncertutil -encode C:\\temp\\mylab-root-ca.cer C:\\temp\\mylab-root-ca.pem")
step("AD>ORA", 12, "Securely ship ora01.keytab and mylab-root-ca.pem to ora01. Delete from C:\\temp on the DC after transfer.")

heading("Phase 2 — Oracle server setup (on ora01)", 1)
step("ORA", 13, "Ensure /etc/hosts has both addresses (ad1 and ora01).")
step("ORA", 14, "Trust the Root CA at the OS level.")
code("sudo cp mylab-root-ca.pem /etc/pki/ca-trust/source/anchors/\nsudo update-ca-trust extract")
success("openssl s_client -connect ad1.mylab.local:636 returns Verify return code: 0 (ok).")
step("ORA", 15, "Install the keytab with the correct permissions.")
code("sudo install -m 0640 -o oracle -g oinstall ora01.keytab /etc/oracle/keytabs/ora01.keytab\nsudo -u oracle klist -kte /etc/oracle/keytabs/ora01.keytab")
success("One entry, oracle/ora01.mylab.local@MYLAB.LOCAL, aes256-cts-hmac-sha1-96, KVNO >= 2.")
step("ORA", 16, "Write /etc/krb5.conf with realm MYLAB.LOCAL and KDC ad1.mylab.local, then run a kinit smoke-test.")
step("ORA", 17, "Configure sqlnet.ora (full content in RECIPE.md). Do NOT set SQLNET.KERBEROS5_CC_NAME with %{uid} — Oracle does not substitute it.")
code("SQLNET.AUTHENTICATION_SERVICES = (BEQ, KERBEROS5)\nSQLNET.KERBEROS5_CONF = /etc/krb5.conf\nSQLNET.KERBEROS5_KEYTAB = /etc/oracle/keytabs/ora01.keytab\nSQLNET.FALLBACK_AUTHENTICATION = FALSE\nWALLET_LOCATION = (SOURCE=(METHOD=FILE)(METHOD_DATA=(DIRECTORY=/u01/app/oracle/cmu/wallet)))")
step("ORA", 18, "Build the Oracle wallet: orapki wallet create, add the Root CA as trusted_cert, add ORACLE.SECURITY.USERNAME/DN/PASSWORD entries for svc-ora-ldap, then enable auto_login. Full recipe in BUILD-STEPS.md step 18.")
step("ORA", 19, "Validate the wallet end-to-end by running a DBMS_LDAP test from PL/SQL (init -> open_ssl -> simple_bind_s).")
note("If this test fails, fix the wallet before continuing — the sync package will fail the same way.")
step("ORA", 20, "Clone the repo onto ora01. Create .env from .env.example and set LDAP_BIND_PWD to the real svc-ora-ldap password (chmod 600).")
step("ORA", 21, "Install the sync package + scheduler + trigger via the wrapper.")
code("sudo -u oracle bash scripts/oracle/run-ad-sync-install.sh")
success("The final state report shows ALICE@MYLAB.LOCAL, BOB@MYLAB.LOCAL, CAROL@MYLAB.LOCAL (EXTERNAL) and matching ORA_*_ROLE grants.")
step("ORA", 22, "Check the sync log for clean execution (no ERROR rows).")
code("SQL> ALTER SESSION SET CONTAINER = orclpdb1;\nSQL> SELECT TO_CHAR(ts,'HH24:MI:SS.FF3'), lvl, msg FROM ad_sync.ad_sync_log ORDER BY ts;")

heading("Phase 3 — Windows client setup (on wks01)", 1)
step("WKS", 23, "Machine domain-joined to MYLAB.LOCAL. (Non-domain laptop: install MIT Kerberos for Windows + use a file ccache.)")
step("WKS", 24, "Import the Root CA cert into Cert:\\LocalMachine\\Root.")
step("WKS", 25, "Enable the allowtgtsessionkey registry key so the JVM can read the TGT from the LSA cache. Reboot afterwards.")
code("reg add HKLM\\SYSTEM\\CurrentControlSet\\Control\\Lsa\\Kerberos\\Parameters /v allowtgtsessionkey /t REG_DWORD /d 1 /f")
step("WKS", 26, "Install DBeaver Community 25.x: winget install dbeaver.dbeaver")
step("WKS", 27, "Copy the four Oracle JDBC jars: ojdbc8.jar (23.3.0.23.09), oraclepki.jar, osdt_core.jar, osdt_cert.jar (21.9.0.0).")
step("WKS", 28, "Edit dbeaver.ini — paste the JVM args under -vmargs. You MUST include the full --add-opens set or Kerberos fails. Restart DBeaver.")
step("WKS", 29, "Create C:\\ProgramData\\MIT\\Kerberos5\\krb5.ini. Critical line: forwardable = false — the fix for the JDBC \"EncryptionKey: Key bytes cannot be null\" bug.")
step("WKS", 30, "In DBeaver create an Oracle connection: Host=ora01.mylab.local, Port=1521, Service Name=orclpdb1, leave username and password blank.")
step("WKS", 31, "On the Driver properties tab set four properties:")
code("oracle.net.authentication_services         = KERBEROS5   (no parens!)\noracle.net.kerberos5_mutual_authentication = false\noracle.net.kerberos5_cc_name               = C:/Users/<user>/krb5cc\noracle.net.kerberos5_conf                  = C:/ProgramData/MIT/Kerberos5/krb5.ini")
step("WKS", 32, "Confirm data-sources.json has \"auth-model\": \"oracle_native\" — NOT \"oracle_os\".")

heading("Phase 4 — End-to-end validation", 1)
para("Run the queries below in DBeaver on wks01, signed in as MYLAB\\alice.", bold=True)
step("WKS", 33, "Open DBeaver, double-click the ORCLPDB1 connection. No password prompt.")
step("WKS", 34, "Run: SELECT USER FROM DUAL;")
success("Returns ALICE@MYLAB.LOCAL.")
step("WKS", 35, "Check auth method: SELECT SYS_CONTEXT('USERENV','AUTHENTICATION_METHOD') FROM DUAL;")
success("Returns KERBEROS | ALICE@MYLAB.LOCAL | alice@MYLAB.LOCAL.")
step("WKS", 36, "Check roles: SELECT ROLE FROM SESSION_ROLES;")
success("Includes ORA_READERS_ROLE (or ORA_WRITERS_ROLE for bob).")
step("WKS", 37, "Run a query that exercises the role: SELECT COUNT(*) FROM ALL_TABLES;")
success("Returns a number > 0. If all four checks pass, the build is complete.")

heading("Phase 5 — Day-2 operations", 1)
step("AD", 38, "Add a user: Add-ADGroupMember -Identity oracle-readers -Members <sam>. Within 10 minutes the scheduler creates the user; or the logon trigger does it instantly on their next connect.")
step("AD", 39, "Move a user between groups — Remove-ADGroupMember from the old, Add-ADGroupMember to the new. Next sync flips the role grants.")
step("ORA", 40, "Force a sync now: BEGIN ad_sync.ad_sync.run; END;")
step("ORA", 41, "Inspect what the last sync did — SELECT ... FROM ad_sync.ad_sync_log.")
step("AD+ORA", 42, "Rotate the svc-ora01 keytab — see docs/10. In short: re-run ktpass, ship to ora01, lsnrctl reload.")
step("AD+ORA", 43, "Rotate the svc-ora-ldap password — Set-ADAccountPassword in AD, then mkstore -modifyEntry ORACLE.SECURITY.PASSWORD on ora01 + lsnrctl reload.")

heading("When things break", 1)
para("Full error-code reference is in troubleshooting.md. Common starting points:")
for ekey, en in [
 ("Clock skew / KRB_AP_ERR_SKEW", "clock skew — resync clocks (NTP/chrony/w32time)."),
 ("KDC_ERR_S_PRINCIPAL_UNKNOWN", "SPN missing or duplicated — setspn -Q / setspn -D."),
 ("KRB_AP_ERR_MODIFIED", "keytab vs AD password mismatch — regenerate the keytab."),
 ("EncryptionKey: Key bytes cannot be null", "set forwardable = false in krb5.ini."),
 ("ORA-24247", "missing network ACL — DBMS_NETWORK_ACL_ADMIN.append_host_ace."),
 ("ORA-28030", "CMU broken on 19c — use the ad_sync pattern (docs/16, docs/17)."),
]:
    p = doc.add_paragraph()
    rc = p.add_run(ekey + "  "); rc.font.name = MONO_FONT; rc.font.size = Pt(9.5); rc.bold = True
    rh = p.add_run("— " + en); rh.font.name = BODY_FONT; rh.font.size = Pt(10.5)

base = r"c:\Users\eranmar\Documents\oracle_Ad_kerberos"
final = os.path.join(base, "Setup-Guide.docx")
tmp = os.path.join(base, "Setup-Guide-NEW.docx")
doc.save(tmp)
try:
    os.replace(tmp, final)
    print("SAVED bytes:", os.path.getsize(final), "-> Setup-Guide.docx")
except PermissionError:
    print("LOCKED: close Setup-Guide.docx in Word. Left as Setup-Guide-NEW.docx ("
          + str(os.path.getsize(tmp)) + " bytes)")
