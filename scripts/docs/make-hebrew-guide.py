# -*- coding: utf-8 -*-
"""Generate a fully-RTL Hebrew .docx build guide.

Key correctness point: OOXML enforces child-element order inside <w:pPr>
and <w:rPr>. python-docx's OxmlElement.append() does NOT, so a misplaced
<w:bidi>/<w:rtl> is silently ignored by Word and the paragraph renders LTR.
This script inserts every RTL element at its schema-correct position, and
also sets RTL at docDefaults + every style so all content inherits it.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEB_FONT = "David"
MONO_FONT = "Consolas"

# --- OOXML schema child order (subset we touch) -----------------------------
PPR_ORDER = ["pStyle","keepNext","keepLines","pageBreakBefore","framePr",
 "widowControl","numPr","suppressLineNumbers","pBdr","shd","tabs",
 "suppressAutoHyphens","kinsoku","wordWrap","overflowPunct","topLinePunct",
 "autoSpaceDE","autoSpaceDN","bidi","adjustRightInd","snapToGrid","spacing",
 "ind","contextualSpacing","mirrorIndents","suppressOverlap","jc",
 "textDirection","textAlignment","textboxTightWrap","outlineLvl","divId",
 "cnfStyle","rPr","sectPr","pPrChange"]
RPR_ORDER = ["rStyle","rFonts","b","bCs","i","iCs","caps","smallCaps","strike",
 "dstrike","outline","shadow","emboss","imprint","noProof","snapToGrid",
 "vanish","webHidden","color","spacing","w","kern","position","sz","szCs",
 "highlight","u","effect","bdr","shd","fitText","vertAlign","rtl","cs","em",
 "lang","eastAsianLayout","specVanish","oMath"]

def _ins(parent, child, order):
    tag = child.tag.split("}")[-1]
    idx = order.index(tag)
    for ex in parent:
        et = ex.tag.split("}")[-1]
        if et in order and order.index(et) > idx:
            ex.addprevious(child); return
    parent.append(child)

def _mk(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e

def ppr_rtl(pPr, jc="right"):
    _ins(pPr, _mk("w:bidi"), PPR_ORDER)
    _ins(pPr, _mk("w:jc", **{"w:val": jc}), PPR_ORDER)

def rpr_rtl(rPr, font=HEB_FONT):
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = _mk("w:rFonts"); _ins(rPr, rf, RPR_ORDER)
    rf.set(qn("w:ascii"), font); rf.set(qn("w:hAnsi"), font); rf.set(qn("w:cs"), font)
    _ins(rPr, _mk("w:rtl"), RPR_ORDER)

doc = Document()

# --- docDefaults: make the whole document RTL by default --------------------
styles_el = doc.styles.element
docDefaults = styles_el.find(qn("w:docDefaults"))
pPrDefault = docDefaults.find(qn("w:pPrDefault"))
if pPrDefault is None:
    pPrDefault = _mk("w:pPrDefault"); docDefaults.insert(0, pPrDefault)
ppd = pPrDefault.find(qn("w:pPr"))
if ppd is None:
    ppd = _mk("w:pPr"); pPrDefault.append(ppd)
ppr_rtl(ppd)
rPrDefault = docDefaults.find(qn("w:rPrDefault"))
rpd = rPrDefault.find(qn("w:rPr"))
if rpd is None:
    rpd = _mk("w:rPr"); rPrDefault.append(rpd)
rpr_rtl(rpd)

# --- apply RTL to Normal + all heading/title styles -------------------------
for st in doc.styles:
    try:
        el = st.element
    except Exception:
        continue
    nm = (st.name or "").lower()
    if nm == "normal" or nm.startswith("heading") or nm == "title":
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            pPr = _mk("w:pPr"); el.insert(0, pPr)
        ppr_rtl(pPr)
        rPr = el.find(qn("w:rPr"))
        if rPr is None:
            rPr = _mk("w:rPr"); el.append(rPr)
        rpr_rtl(rPr)

normal = doc.styles["Normal"]
normal.font.name = HEB_FONT
normal.font.size = Pt(11)

# --- section RTL ------------------------------------------------------------
sectPr = doc.sections[0]._sectPr
# bidi goes near end of sectPr; order is lenient enough, but place before pgNumType/cols issues:
sectPr.append(_mk("w:bidi"))

# --- helpers ----------------------------------------------------------------
def _finalize(p, runs_font=HEB_FONT, jc="right"):
    # NOTE: do NOT also call p.alignment = ... — python-docx would rewrite the
    # w:jc we set here (clobbering "both" back to "right").
    pPr = p._p.get_or_add_pPr()
    ppr_rtl(pPr, jc=jc)
    for r in p.runs:
        rPr = r._element.get_or_add_rPr()
        rpr_rtl(rPr, runs_font)

def heading(text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = HEB_FONT
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)
    _finalize(p)
    return p

def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold
    r.font.name = HEB_FONT; r.font.size = Pt(11)
    _finalize(p, jc="both")
    return p

def step(tag, num, text):
    p = doc.add_paragraph()
    r0 = p.add_run(f"{num}. "); r0.bold = True; r0.font.name = HEB_FONT
    rt = p.add_run(f"[{tag}] "); rt.bold = True; rt.font.name = MONO_FONT
    rt.font.color.rgb = RGBColor(0xB0, 0x30, 0x10)
    rb = p.add_run(text); rb.font.name = HEB_FONT; rb.font.size = Pt(11)
    _finalize(p, jc="both")
    p.paragraph_format.space_after = Pt(4)
    return p

def success(text):
    p = doc.add_paragraph()
    r0 = p.add_run("✔ הצלחה: ")
    r0.bold = True; r0.font.name = HEB_FONT
    r0.font.color.rgb = RGBColor(0x1B, 0x7A, 0x1B)
    r1 = p.add_run(text); r1.font.name = HEB_FONT; r1.font.size = Pt(10.5)
    _finalize(p, jc="both")
    p.paragraph_format.space_after = Pt(8)
    return p

def note(text):
    p = doc.add_paragraph()
    r0 = p.add_run("⚠ הערה: ")
    r0.bold = True; r0.font.name = HEB_FONT
    r0.font.color.rgb = RGBColor(0xB0, 0x60, 0x00)
    r1 = p.add_run(text); r1.font.name = HEB_FONT; r1.italic = True; r1.font.size = Pt(10.5)
    _finalize(p, jc="both")
    return p

def code(text):
    """LTR monospace block, left-aligned, shaded. Commands stay LTR."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.3); pf.right_indent = Inches(0.3)
    pf.space_before = Pt(2); pf.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    # explicit LTR: no bidi; jc=left; shading
    _ins(pPr, _mk("w:shd", **{"w:val": "clear", "w:fill": "F2F2F2"}), PPR_ORDER)
    _ins(pPr, _mk("w:jc", **{"w:val": "left"}), PPR_ORDER)
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line); r.font.name = MONO_FONT; r.font.size = Pt(9.5)
    return p

# ===========================================================================
heading("מדריך הקמה: אימות Active Directory מול Oracle באמצעות Kerberos", 0)
para("מסמך זה מתאר שלב-אחר-שלב מה צריך לעשות כדי שמשתמש דומיין (Active Directory) יוכל להתחבר למסד הנתונים Oracle באמצעות Kerberos, דרך DBeaver ב-Windows, ללא סיסמה מקומית ב-Oracle. כל שלב מסומן בתגית המכונה: [AD] = בקר הדומיין, [ORA] = שרת Oracle, [WKS] = תחנת העבודה.")
note("כל הפקודות, שמות המוצרים ונתיבי הקבצים באנגלית ומיושרים לשמאל — העתק כמו שהם. ערכי המעבדה: MYLAB.LOCAL, ad1.mylab.local, ora01.mylab.local, wks01.mylab.local, קבוצות oracle-readers / oracle-writers, משתמשי בדיקה alice / bob / carol.")

heading("שלב 0 — דרישות מקדימות", 1)
step("AD", 1, "בקר דומיין Windows Server 2022, יער mylab.local, תפקיד DNS מותקן, רשומות A עבור ad1, ora01, wks01.")
step("AD", 2, "שרותי אישורים (AD CS) מותקנים — Enterprise Root CA בשם mylab-root-ca. ודא שהבקר קיבל תעודת LDAPS (subject CN=ad1.mylab.local).")
step("ORA", 3, "שרת RHEL 9 עם Oracle 19c (19.30+), CDB בשם ORCLCDB פתוח, PDB בשם ORCLPDB1 פתוח READ WRITE, Listener על פורט 1521.")
step("ALL", 4, "כל שלוש המכונות באותה רשת. בדוק נגישות לפורטים 88, 389, 636, 464.")
step("ALL", 5, "פער שעון בין כל זוג מכונות קטן מ-300 שניות (חלון Kerberos). ora01 מסתנכרן עם ad1 דרך chronyd.")

heading("שלב 1 — הגדרת Active Directory (על ad1)", 1)
step("AD", 6, "צור חשבון שירות svc-ora01 עם סיסמה זמנית (placeholder בלבד — AD לא מאפשר חשבון פעיל ללא סיסמה; שלב 7 יאפס אותה לערך האמיתי). אם החשבון כבר קיים — דלג על New-ADUser והפעל רק Enable-ADAccount ו-Set-ADUser.")
code("$throwaway = ConvertTo-SecureString ([guid]::NewGuid().ToString()+'Aa1!') -AsPlainText -Force\nNew-ADUser -Name svc-ora01 -SamAccountName svc-ora01 `\n  -UserPrincipalName svc-ora01@MYLAB.LOCAL -AccountPassword $throwaway `\n  -Enabled $true -PasswordNeverExpires $true -CannotChangePassword $true `\n  -Path \"CN=Users,DC=mylab,DC=local\"\nSet-ADUser -Identity svc-ora01 -KerberosEncryptionType \"AES256\"")
success("Get-ADUser svc-ora01 מחזיר תוצאה; msDS-SupportedEncryptionTypes שווה 16.")
step("AD", 7, "רשום את ה-SPN וצור את ה-keytab בפעולה אחת עם ktpass. השתמש ב-‎-pass *‎ כך ש-ktpass יבקש את הסיסמה באופן מוסתר. זו הסיסמה המחייבת — שמור אותה בכספת הסיסמאות. ktpass תמיד מאפס את הסיסמה ומעלה KVNO — לכן הרצת הפקודה הזו גם ביצירה וגם בסיבוב סיסמאות זהה.")
code("ktpass -princ oracle/ora01.mylab.local@MYLAB.LOCAL `\n       -mapuser MYLAB\\svc-ora01 `\n       -pass * `\n       -crypto AES256-SHA1 -ptype KRB5_NT_PRINCIPAL `\n       -out C:\\temp\\ora01.keytab")
note("שני הדגלים שטועים בהם: -crypto AES256-SHA1 = הפק רק מפתח AES256 (etype 18) ל-keytab; חייב להתאים ל-msDS-SupportedEncryptionTypes של החשבון (=16), ל-sqlnet.ora ול-krb5.ini בלקוח; לעולם אל תשתמש ב-All (מכניס מפתח RC4 חלש). -ptype KRB5_NT_PRINCIPAL = סוג-שם עיקרון תקני ש-Oracle/MIT/Java מצפים לו; סוגים אחרים גורמים לכשל אימות. טבלה מלאה + אימות AES: docs/19-ad-admin-runbook.md סעיף A2. לאחר ההפקה ודא ש-klist -kte מציג etype 18 ו-Get-ADUser svc-ora01 -Properties msDS-SupportedEncryptionTypes מחזיר 16.")
note("אם setspn -Q oracle/ora01.mylab.local מחזיר יותר מחשבון אחד או חשבון שאינו svc-ora01 — הסר את ה-SPN מכל בעלים שגוי: setspn -D oracle/ora01.mylab.local <חשבון-שגוי>, והרץ שוב ktpass.")
success("setspn -Q oracle/ora01.mylab.local מחזיר בדיוק חשבון אחד (CN=svc-ora01). setspn -X ללא כפילויות.")
step("AD", 8, "צור חשבון svc-ora-ldap (Oracle מתחבר איתו ל-AD ב-LDAPS לקריאת חברות בקבוצות). הרשאות Domain Users מספיקות. בניגוד ל-svc-ora01, לחשבון זה אין keytab — הסיסמה נטענת ישירות ל-Wallet, ולכן הסיסמה כאן היא המחייבת. אם החשבון כבר קיים — אל תשנה את הסיסמה (זה ישבור את הסנכרון).")
code("if (-not (Get-ADUser -Filter \"SamAccountName -eq 'svc-ora-ldap'\" `\n          -ErrorAction SilentlyContinue)) {\n  $pw = Read-Host -AsSecureString \"NEW password for svc-ora-ldap\"\n  New-ADUser -Name svc-ora-ldap -SamAccountName svc-ora-ldap `\n    -UserPrincipalName svc-ora-ldap@MYLAB.LOCAL -AccountPassword $pw `\n    -Enabled $true -PasswordNeverExpires $true -CannotChangePassword $true `\n    -Path \"CN=Users,DC=mylab,DC=local\"\n} else { Write-Host \"exists - leave password (rotate via Part B3)\" }")
success("Get-ADUser svc-ora-ldap מחזיר חשבון מופעל.")
step("AD", 9, "צור את ה-OU והקבוצות: OU=Groups,DC=mylab,DC=local המכיל oracle-readers ו-oracle-writers (Global Security). אידמפוטנטי — בטוח להרצה חוזרת.")
code("if (-not (Get-ADOrganizationalUnit -Filter 'Name -eq \"Groups\"' `\n          -SearchBase \"DC=mylab,DC=local\" -ErrorAction SilentlyContinue)) {\n  New-ADOrganizationalUnit -Name \"Groups\" -Path \"DC=mylab,DC=local\"\n}\nNew-ADGroup -Name oracle-readers -GroupScope Global -GroupCategory Security `\n  -Path \"OU=Groups,DC=mylab,DC=local\"\nNew-ADGroup -Name oracle-writers -GroupScope Global -GroupCategory Security `\n  -Path \"OU=Groups,DC=mylab,DC=local\"")
step("AD", 10, "הוסף משתמשי בדיקה לקבוצות (alice ו-carol ל-oracle-readers, bob ל-oracle-writers).")
code("Add-ADGroupMember -Identity oracle-readers -Members alice, carol\nAdd-ADGroupMember -Identity oracle-writers -Members bob")
step("AD", 11, "ייצא את תעודת ה-Root CA כך ש-Wallet של Oracle יוכל לבטוח בשרשרת ה-LDAPS.")
code("(Get-CACertificate).RawData | Set-Content -Encoding Byte C:\\temp\\mylab-root-ca.cer\ncertutil -encode C:\\temp\\mylab-root-ca.cer C:\\temp\\mylab-root-ca.pem")
step("AD>ORA", 12, "העבר באופן מאובטח את ora01.keytab ו-mylab-root-ca.pem ל-ora01. מחק מ-C:\\temp ב-DC לאחר ההעברה.")

heading("שלב 2 — הגדרת שרת Oracle (על ora01)", 1)
step("ORA", 13, "ודא ש-/etc/hosts מכיל את שתי הכתובות (ad1 ו-ora01).")
step("ORA", 14, "בטח את ה-Root CA ברמת מערכת ההפעלה.")
code("sudo cp mylab-root-ca.pem /etc/pki/ca-trust/source/anchors/\nsudo update-ca-trust extract")
success("openssl s_client -connect ad1.mylab.local:636 מחזיר Verify return code: 0 (ok).")
step("ORA", 15, "התקן את ה-keytab עם ההרשאות הנכונות.")
code("sudo install -m 0640 -o oracle -g oinstall ora01.keytab /etc/oracle/keytabs/ora01.keytab\nsudo -u oracle klist -kte /etc/oracle/keytabs/ora01.keytab")
success("רשומה אחת, oracle/ora01.mylab.local@MYLAB.LOCAL, aes256-cts-hmac-sha1-96, KVNO ≥ 2.")
step("ORA", 16, "כתוב /etc/krb5.conf עם היער MYLAB.LOCAL וה-KDC ad1.mylab.local, ובצע בדיקת kinit.")
step("ORA", 17, "הגדר את sqlnet.ora (תוכן מלא ב-RECIPE.md). אל תגדיר SQLNET.KERBEROS5_CC_NAME עם %{uid}.")
code("SQLNET.AUTHENTICATION_SERVICES = (BEQ, KERBEROS5)\nSQLNET.KERBEROS5_CONF = /etc/krb5.conf\nSQLNET.KERBEROS5_KEYTAB = /etc/oracle/keytabs/ora01.keytab\nSQLNET.FALLBACK_AUTHENTICATION = FALSE\nWALLET_LOCATION = (SOURCE=(METHOD=FILE)(METHOD_DATA=(DIRECTORY=/u01/app/oracle/cmu/wallet)))")
step("ORA", 18, "בנה את ה-Wallet: orapki wallet create, הוספת trusted_cert של Root CA, ורשומות ORACLE.SECURITY.USERNAME/DN/PASSWORD עבור svc-ora-ldap, ולבסוף auto_login. פרטים ב-BUILD-STEPS.md שלב 18.")
step("ORA", 19, "ודא את ה-Wallet מקצה-לקצה על ידי הרצת DBMS_LDAP מתוך PL/SQL (init → open_ssl → simple_bind_s).")
note("אם הבדיקה נכשלת — תקן את ה-Wallet לפני שממשיכים. חבילת הסנכרון תיכשל באותו אופן.")
step("ORA", 20, "שבט את ה-repo. צור .env מ-.env.example ומלא LDAP_BIND_PWD עם הסיסמה האמיתית של svc-ora-ldap (הרשאות 0600).")
step("ORA", 21, "התקן את חבילת הסנכרון + מתזמן + טריגר על ידי הרצת ה-wrapper.")
code("sudo -u oracle bash scripts/oracle/run-ad-sync-install.sh")
success("דוח המצב מציג את ALICE@MYLAB.LOCAL, BOB@MYLAB.LOCAL, CAROL@MYLAB.LOCAL (EXTERNAL) והרשאות ORA_*_ROLE.")
step("ORA", 22, "בדוק את טבלת ה-log של הסנכרון לוודא שאין שורות ERROR.")
code("SQL> ALTER SESSION SET CONTAINER = orclpdb1;\nSQL> SELECT TO_CHAR(ts,'HH24:MI:SS.FF3'), lvl, msg FROM ad_sync.ad_sync_log ORDER BY ts;")

heading("שלב 3 — הגדרת הלקוח (על wks01)", 1)
step("WKS", 23, "המחשב מחובר לדומיין MYLAB.LOCAL. (מחשב לא-מדומיין — התקן MIT Kerberos for Windows + file ccache).")
step("WKS", 24, "ייבא את תעודת ה-Root CA ל-Cert:\\LocalMachine\\Root.")
step("WKS", 25, "הפעל את מפתח הרגיסטרי allowtgtsessionkey כך ש-JVM יכול לקרוא את כרטיס ה-TGT. אתחל מחדש לאחר מכן.")
code("reg add HKLM\\SYSTEM\\CurrentControlSet\\Control\\Lsa\\Kerberos\\Parameters /v allowtgtsessionkey /t REG_DWORD /d 1 /f")
step("WKS", 26, "התקן DBeaver Community 25.x: winget install dbeaver.dbeaver")
step("WKS", 27, "העתק את ארבעת קבצי ה-jar: ojdbc8.jar (23.3.0.23.09), oraclepki.jar, osdt_core.jar, osdt_cert.jar (21.9.0.0).")
step("WKS", 28, "ערוך את dbeaver.ini — הדבק את ארגומנטי ה-JVM מתחת ל-‎-vmargs‎. חובה לכלול את כל סט ה-‎--add-opens‎ — בלעדיו Kerberos נכשל. אתחל מחדש את DBeaver.")
step("WKS", 29, "צור את C:\\ProgramData\\MIT\\Kerberos5\\krb5.ini. שורה קריטית: forwardable = false — זה התיקון לבאג EncryptionKey: Key bytes cannot be null.")
step("WKS", 30, "ב-DBeaver צור חיבור Oracle: Host=ora01.mylab.local, Port=1521, Service Name=orclpdb1, שם משתמש וסיסמה — השאר ריק.")
step("WKS", 31, "בלשונית Driver properties הגדר ארבעה מאפיינים:")
code("oracle.net.authentication_services         = KERBEROS5   (no parens!)\noracle.net.kerberos5_mutual_authentication = false\noracle.net.kerberos5_cc_name               = C:/Users/<user>/krb5cc\noracle.net.kerberos5_conf                  = C:/ProgramData/MIT/Kerberos5/krb5.ini")
step("WKS", 32, "ודא ש-data-sources.json מכיל \"auth-model\": \"oracle_native\" — ולא \"oracle_os\".")

heading("שלב 4 — אימות מקצה-לקצה", 1)
para("הרץ את השאילתות הבאות ב-DBeaver על wks01, כשאתה מחובר כ-MYLAB\\alice.", bold=True)
step("WKS", 33, "פתח את DBeaver, לחץ פעמיים על חיבור ORCLPDB1. ללא בקשת סיסמה.")
step("WKS", 34, "הרץ: SELECT USER FROM DUAL;")
success("מחזיר ALICE@MYLAB.LOCAL.")
step("WKS", 35, "בדוק שיטת אימות: SELECT SYS_CONTEXT('USERENV','AUTHENTICATION_METHOD') FROM DUAL;")
success("מחזיר KERBEROS | ALICE@MYLAB.LOCAL | alice@MYLAB.LOCAL.")
step("WKS", 36, "בדוק תפקידים: SELECT ROLE FROM SESSION_ROLES;")
success("כולל ORA_READERS_ROLE (או ORA_WRITERS_ROLE עבור bob).")
step("WKS", 37, "הרץ שאילתה שממשת את התפקיד: SELECT COUNT(*) FROM ALL_TABLES;")
success("מחזיר מספר > 0. אם ארבעת המבחנים עוברים — ההקמה הושלמה.")

heading("שלב 5 — תפעול שוטף (Day-2)", 1)
step("AD", 38, "הוספת משתמש: Add-ADGroupMember -Identity oracle-readers -Members <sam>. תוך 10 דקות המתזמן יוצר את המשתמש; או הטריגר מידית בהתחברות הבאה.")
step("AD", 39, "העברת משתמש בין קבוצות — Remove-ADGroupMember מהישנה, Add-ADGroupMember לחדשה.")
step("ORA", 40, "הרץ סנכרון מידי: BEGIN ad_sync.ad_sync.run; END;")
step("ORA", 41, "בדוק מה הסנכרון האחרון עשה — SELECT ... FROM ad_sync.ad_sync_log.")
step("AD+ORA", 42, "סיבוב keytab של svc-ora01 — ראה docs/10. בקצרה: הרץ ktpass מחדש, העבר ל-ora01, lsnrctl reload.")
step("AD+ORA", 43, "סיבוב סיסמת svc-ora-ldap — Set-ADAccountPassword ב-AD, אז mkstore -modifyEntry ORACLE.SECURITY.PASSWORD ב-ora01 + lsnrctl reload.")

heading("כשמשהו נשבר", 1)
para("הפניות מלאות לפי קוד שגיאה ב-troubleshooting.md. נקודות התחלה נפוצות:")
for ekey, heb in [
 ("Clock skew / KRB_AP_ERR_SKEW", "פער שעון — סנכרן שעונים (NTP/chrony/w32time)."),
 ("KDC_ERR_S_PRINCIPAL_UNKNOWN", "SPN חסר או כפול — setspn -Q / setspn -D."),
 ("KRB_AP_ERR_MODIFIED", "אי-התאמה בין keytab לסיסמת AD — צור keytab מחדש."),
 ("EncryptionKey: Key bytes cannot be null", "הגדר forwardable = false ב-krb5.ini."),
 ("ORA-24247", "חסר אישור ACL לרשת — DBMS_NETWORK_ACL_ADMIN.append_host_ace."),
 ("ORA-28030", "CMU שבור ב-19c — השתמש בתבנית ad_sync (docs/16, docs/17)."),
]:
    p = doc.add_paragraph()
    rc = p.add_run(ekey + "  "); rc.font.name = MONO_FONT; rc.font.size = Pt(9.5); rc.bold = True
    rh = p.add_run("— " + heb); rh.font.name = HEB_FONT; rh.font.size = Pt(10.5)
    _finalize(p)

base = r"c:\Users\eranmar\Documents\oracle_Ad_kerberos"
final = os.path.join(base, "מדריך-הקמה.docx")
tmp   = os.path.join(base, "מדריך-הקמה-NEW.docx")
doc.save(tmp)
try:
    os.replace(tmp, final)
    print("SAVED bytes:", os.path.getsize(final), "-> מדריך-הקמה.docx")
except PermissionError:
    print("LOCKED: original open in Word. New file left as מדריך-הקמה-NEW.docx ("
          + str(os.path.getsize(tmp)) + " bytes)")
